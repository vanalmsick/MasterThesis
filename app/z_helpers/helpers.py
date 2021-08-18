import sys, os, pathlib, warnings, datetime, mlflow, json, tabulate, hashlib
import datetime as dt
import tables as tb
import pandas as pd
import numpy as np
import psycopg2, psycopg2.extras
import configparser
from silx.io import dictdump
from tqdm import tqdm
import multiprocessing
import eikon as ek


######################################### GENERAL #########################################


def change_workdir():
    new_wkdir = os.path.dirname(os.path.dirname(__file__))
    os.chdir(new_wkdir)
    print('Working dir changed to:', os.getcwd())



def convenience_settings():
    pd.set_option('display.max_rows', 50)
    pd.set_option('display.max_columns', 12)
    pd.set_option('display.width', 300)
    print('Custom parameters set')
    change_workdir()


def get_project_directories(key=None, **kwargs):
    helpers_path = os.path.dirname(__file__)
    app_path = os.path.dirname(helpers_path)
    project_path = os.path.dirname(app_path)
    data_path = os.path.join(project_path, 'data')
    cache_path = os.path.join(project_path, 'cache')
    credentials_path = os.path.join(project_path, 'credentials')
    working_dir = os.getcwd()
    tensorboard_logs = '/Users/vanalmsick/opt/anaconda3/envs/dev/lib/python3.8/site-packages/tensorboard/logs'

    dir_dict = {'project_dir': project_path, 'app_dir': app_path, 'data_dir': data_path, 'cred_dir':credentials_path, 'helpers_dir': helpers_path, 'working_dir': working_dir, 'tensorboard_logs':tensorboard_logs, 'cache_dir':cache_path}

    if key is None:
        return dir_dict
    else:
        return dir_dict[key]


def get_curr_time():
    return dt.datetime.now().strftime("%Y.%m.%d.%H.%M.%S")

def get_start_time():
    return get_curr_time()


def get_credentials(credential=None, cred_dir=None, dir_dict=None, **kwargs):
    if cred_dir is None and dir_dict is None:
        cred_dir = get_project_directories(key='cred_dir')
    elif cred_dir is None and dir_dict is not None:
        cred_dir = dir_dict['cred_dir']

    file_list = [f for f in os.listdir(cred_dir) if os.path.isfile(os.path.join(cred_dir, f))]
    creds = {}

    for file in file_list:
        if file[-4:] == '.ini':
            config = configparser.ConfigParser()
            config.read(os.path.join(cred_dir, file))
            conf_dict = config._sections
            if len(conf_dict.keys()) == 1:
                key = list(conf_dict.keys())[0]
                creds[key] = conf_dict[key]
            else:
                creds[file.split('.')[0]] = conf_dict
        elif file[-4:] == '.key':
            with open(os.path.join(cred_dir, file), 'r') as f:
                creds[file.split('.')[0]] = f.read()

    if credential is None:
        return creds
    else:
        return creds[credential]



def df_iloc_several_ranges(df, *tuple_pairs):
    """
    Input pairs of tuples for index slicing
    """
    len_df = len(df)

    # Create an array with values to use as an index
    num_range = np.zeros(shape=(len_df,), dtype=bool)

    # Update
    for (start, end) in tuple_pairs:
        if start is None:
            start = 0
        if end is None:
            end = len_df
        num_range[start:end] = True

    return df.iloc[num_range]



def data_hash(self, *args):
    str_args = (str(args)[1:-1]).replace("'", "").replace(", ", "/")
    hash = hashlib.shake_256(str_args.encode()).hexdigest(5)
    return hash




def highlight_dataframe_differences(df1, df2):
    """Identify differences between two pandas DataFrames"""
    assert (df1.columns == df2.columns).all(), \
        "DataFrame column names are different"
    if any(df1.dtypes != df2.dtypes):
        "Data Types are different, trying to convert"
        df2 = df2.astype(df1.dtypes)
    if df1.equals(df2):
        return None
    else:
        # need to account for np.nan != np.nan returning True
        diff_mask = (df1 != df2) & ~(df1.isnull() & df2.isnull())
        ne_stacked = diff_mask.stack()
        changed = ne_stacked[ne_stacked]
        changed.index.names = ['id', 'col']
        difference_locations = np.where(diff_mask)
        changed_from = df1.values[difference_locations]
        changed_to = df2.values[difference_locations]
        return pd.DataFrame({'from': changed_from, 'to': changed_to},
                            index=changed.index)





######################################### SQL/Database #########################################


def postgresql_connect(params_dic):
    """ Connect to the PostgreSQL database server """
    conn = None
    try:
        # connect to the PostgreSQL server
        print('Connecting to the PostgreSQL database...')
        conn = psycopg2.connect(**params_dic)
    except (Exception, psycopg2.DatabaseError) as error:
        print(error)
        sys.exit(1)
    print("Connection successful")
    return conn


def _dtype_mapping():
    return {'object' : 'TEXT',
        'string': 'TEXT',
        'int64' : 'INT', # INT
        'float64' : 'DOUBLE PRECISION',
        'datetime64' : 'DATETIME',
        'bool' : 'BOOLEAN',
        'category' : 'TEXT',
        'timedelta[ns]' : 'TIMESTAMP'}



def _gen_tbl_cols_sql(df):
    dmap = _dtype_mapping()
    sql = ""
    df1 = df.rename(columns = {"" : "nocolname", np.nan : "nocolname"})
    hdrs = df1.dtypes.index
    hdrs_list = [(hdr, str(df1[hdr].dtype)) for hdr in hdrs]
    for i, hl in enumerate(hdrs_list):
        sql += " ,{0} {1}".format(hl[0].lower(), dmap[str(hl[1]).lower()])
    return sql[2:]


def create_sql_tbl(df, conn, tbl_name, schema='public'):
    if check_table_exists(conn=conn, table_name=tbl_name, schema=schema):
        print('Table {}.{} already exists.'.format(schema, tbl_name))
    else:
        tbl_cols_sql = _gen_tbl_cols_sql(df)
        sql = "CREATE TABLE {}.{} ({})".format(schema, tbl_name, tbl_cols_sql)
        cur = conn.cursor()
        cur.execute(sql)
        cur.close()
        conn.commit()


def check_table_exists(conn, table_name, schema='public'):
    sql = """SELECT EXISTS (
               SELECT FROM pg_catalog.pg_class c
               JOIN   pg_catalog.pg_namespace n ON n.oid = c.relnamespace
               WHERE  n.nspname = '{schema}'
               AND    c.relname = '{table}'
               AND    c.relkind = 'r'    -- only tables
               )""".format(schema=schema, table=table_name)
    with conn.cursor() as cur:
        cur.execute(sql)
        answer = cur.fetchall()[0][0]
    return answer


def sql_query(sql, conn):
    return pd.read_sql_query(sql, conn)


def df_insert_sql(conn, df, table, schema='public'):
    """
    Using psycopg2.extras.execute_values() to insert the dataframe
    """
    # Create a list of tupples from the dataframe values
    df = df.replace({pd.NaT: np.nan, '':np.nan})
    for col in (df.select_dtypes(include=[np.datetime64])).columns.tolist():
        df[col] = [d if not pd.isnull(d) else None for d in df[col]]
    warnings.filterwarnings('ignore')
    np_array = df.replace({pd.np.nan: None}).values
    warnings.filterwarnings('default')
    tuples = [tuple(x) for x in np_array]
    # Comma-separated dataframe columns
    cols = ','.join(df.columns.tolist())
    # SQL quert to execute
    query  = "INSERT INTO %s.%s (%s) VALUES %%s" % (schema, table, cols)
    cursor = conn.cursor()
    try:
        psycopg2.extras.execute_values(cursor, query, tuples)
        conn.commit()
    except (Exception, psycopg2.DatabaseError) as error:
        print("Error: %s" % error)
        conn.rollback()
        cursor.close()
        return 1, error
    #print("execute_values() done")
    cursor.close()


def update_request_status(request_id, conn, schema, new_status=1):
    if type(request_id) != list:
        request_id = [request_id]

    with conn.cursor() as cur:
        for id in request_id:
            updated = datetime.datetime.now().isoformat()
            SQL = """UPDATE {schema}.data_request_list SET status = {status}, last_updated = '{updated}' WHERE request_id = {id}""".format(status=new_status, schema=schema, updated=updated, id=id)
            cur.execute(SQL)
    conn.commit()

def submit_error(error_dict, conn, schema):
    request_id = error_dict['request_id']
    error_type = error_dict['error_type']
    error_message = str(error_dict['error_message']).replace("'",'"')
    if error_message[:14] == '[{"code": 416,' and len(error_message) > 41: # shorten if common error message
        error_message = error_message[:40]
    error_comps = str(error_dict['req_instruments']).replace("'",'"')

    if 'timestamp_iso' in error_dict:
        error_time = error_dict['timestamp_iso']
    else:
        error_time = datetime.datetime.now().isoformat()

    with conn.cursor() as cur:
        SQL = """INSERT INTO {schema}.data_request_errors (request_id, error_time, error_type, error_message, comps) VALUES ({request_id}, '{error_time}', '{error_type}', '{error_message}', '{comps}')""".format(schema=schema, request_id=request_id, error_time=error_time, error_type=error_type, error_message=error_message, comps=error_comps)
        cur.execute(SQL)
    conn.commit()


def create_wrds_sql_query(library='comp', table='company', columns='*', conditions=None, custom_condition='', distinct=False, no_observations=-1, offset=0):
    """Libary options: libary=['fundq', 'funda', 'company']"""
    if distinct:
        dis = ' DISTINCT '
    else:
        dis = ''
    if no_observations < 0:
        obsstmt = ''
    else:
        obsstmt = 'LIMIT {}'.format(no_observations)
    if columns is None:
        cols = '*'
    else:
        cols = ','.join(columns)
    if conditions is None:
        cond = ' '
    else:
        cond = 'WHERE'
        if custom_condition != '':
            cond = cond + ' ' + custom_condition + ' AND'
        for key in conditions:
            if type(conditions[key]) == list:
                cond = cond + ' ' + key + ' in (' + str(conditions[key])[1:-1] + ') AND'
            elif type(conditions[key]) == str:
                cond = cond + ' ' + key + " = '" + str(conditions[key]) + "' AND"
            elif type(conditions[key]) == int or type(conditions[key]) == float:
                cond = cond + ' ' + key + " = " + str(conditions[key]) + " AND"
            else:
                raise Exception('Unknown data type of condition.')
        cond = cond[:-3]

    sql = ("SELECT{distinct} {cols} FROM {schema}.{table} {condition} {obsstmt} OFFSET {offset};".format(cols=cols, distinct=dis, schema=library, table=table, condition=cond, obsstmt=obsstmt, offset=offset))
    return sql



def reuters_eikon_data_scraper(instruments: list, fields: list, properties:dict, api_key:str):
    df, err = None, None
    ek.set_app_key(api_key)
    try:
        df, err = ek.get_data(instruments, fields, properties, field_name=True, raw_output=False)
        if err is not None:
            err = {'error_type': 'REUTERS', 'error_message': str(err), 'req_instruments': instruments, 'req_fields': fields, 'req_properties': properties}
        else:
            df.columns = [i.replace('.', '_') for i in df.columns.tolist()]
    except Exception as error:
        err = {'error_type': 'PyREUTERS', 'error_message': str(error).replace('\n', ';').encode(), 'req_instruments': instruments, 'req_fields': fields, 'req_properties':properties}
    finally:
        return df, err


def wrds_compustat_data_scraper(conn, sql):
    df, err = None, None
    try:
        df = conn.raw_sql(sql)
    except Exception as error:
        err = {'error_type': 'PyWRDS', 'error_message': str(error).replace('\n', ';').encode(), 'req_sql': sql}
    finally:
        return df, err





class custom_hdf5:
    def dict_to_hdf5(file, save_dict):
        create_ds_args = {'compression': "gzip", 'fletcher32': True}
        dictdump.dicttoh5(save_dict, file, h5path="/", mode='w', create_dataset_args=create_ds_args)
        #dictdump.dicttoh5(save_dict, file, h5path="/", mode='w')


    def hdf5_to_dict(file, *args):
        hdf_path = '/' + '/'.join(args[:-1])
        out = dictdump.h5todict(file, hdf_path)
        out = out[args[-1]]
        return out

    def hdf5_to_pd(file, *args):
        my_dict = custom_hdf5.hdf5_to_dict(file, *args)
        pd_ser = pd.Series(dict(zip(my_dict[0, :].astype(str), my_dict[1, :].astype(float))))
        return pd_ser

    def pd_series_to_2d_array(pd_series):
        np_array = np.array([pd_series.index.values.astype(str), pd_series.values.astype(float)])
        return np_array

    def get_comp_list(file, norm_key):
        import h5py
        with h5py.File(file, 'r') as f:
            a = list(f.get(norm_key).keys())
        try:
            a.remove('__all__')
        except:
            pass
        return a


from regressors import stats
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
def word_regression_table(model, X, y, output_file, title=None, overwrite_summary={}):
    titel = 'default' if title is None else title
    y_col = y.name
    X_cols = X.columns.tolist()


    def _modified_regressor_summary(clf, X, y, xlabels=None):
        """
        Output summary statistics for a fitted regression model.

        Parameters
        ----------
        clf : sklearn.linear_model
            A scikit-learn linear model classifier with a `predict()` method.
        X : numpy.ndarray
            Training data used to fit the classifier.
        y : numpy.ndarray
            Target training values, of shape = [n_samples].
        xlabels : list, tuple
            The labels for the predictors.
        """
        # Check and/or make xlabels
        ncols = X.shape[1]
        if xlabels is None:
            xlabels = np.array(
                ['x{0}'.format(i) for i in range(1, ncols + 1)], dtype='str')
        elif isinstance(xlabels, (tuple, list)):
            xlabels = np.array(xlabels, dtype='str')
        # Make sure dims of xlabels matches dims of X
        if xlabels.shape[0] != ncols:
            raise AssertionError(
                "Dimension of xlabels {0} does not match "
                "X {1}.".format(xlabels.shape, X.shape))
        # Create data frame of coefficient estimates and associated stats
        coef_df = pd.DataFrame(
            index=['_intercept'] + list(xlabels),
            columns=['Estimate', 'Std. Error', 't value', 'p value']
        )
        coef_df['Estimate'] = np.concatenate(
            (np.round(np.array([clf.intercept_]), 6), np.round((clf.coef_), 6)))
        coef_df['Std. Error'] = np.round(stats.coef_se(clf, X, y), 6)
        coef_df['t value'] = np.round(stats.coef_tval(clf, X, y), 4)
        coef_df['p value'] = np.round(stats.coef_pval(clf, X, y), 6)
        # Create data frame to summarize residuals
        resids = stats.residuals(clf, X, y, r_type='raw')
        resids_df = pd.DataFrame({
            'Min': pd.Series(np.round(resids.min(), 4)),
            '1Q': pd.Series(np.round(np.percentile(resids, q=25), 4)),
            'Median': pd.Series(np.round(np.median(resids), 4)),
            '3Q': pd.Series(np.round(np.percentile(resids, q=75), 4)),
            'Max': pd.Series(np.round(resids.max(), 4)),
        }, columns=['Min', '1Q', 'Median', '3Q', 'Max'])

        return resids_df, coef_df, {'R2': stats.metrics.r2_score(y, clf.predict(X)), 'Adj R2': stats.adj_r2_score(clf, X, y),
                                    'F-statistic': stats.f_stat(clf, X, y)}

    def _round_if_number(s):
        if type(s) == int:
            return s
        try:
            float(s)
            s = format(round(s, 3), '.3f')
            return s
        except ValueError:
            return s


    def _dict_to_np(param):
        rows = int(np.ceil(len(param)/2))
        arr = np.empty((rows, 4), dtype=np.dtype(object))
        for i, key in zip(range(len(param)), param):
            if i >= rows:
                y = i - rows
                x = 2
            else:
                y = i
                x = 0
            arr[y, x] = key
            arr[y, x + 1] = _round_if_number(param[key])
        return arr


    def _df_to_np(df, ignore_idx=False, ignore_head=False):
        cols = df.columns.tolist()
        idxs = df.index.tolist()
        i_idx = 0 if ignore_idx else 1
        i_col = 0 if ignore_head else 1
        arr = np.empty((len(idxs) + i_col, len(cols) + i_idx), dtype=np.dtype(object))
        arr = np.where(arr is None, arr, '')
        # Write index
        if ignore_idx is False:
            for y in range(i_col, len(idxs) + i_col):
                arr[y, 0] = idxs[y - i_col]
        # Write cols
        if ignore_head is False:
            for x in range(i_idx, len(cols) + i_idx):
                arr[0, x] = cols[x - i_idx]
        #  Write values
        df_val = df.values
        for y in range(i_col, len(idxs) + i_col):
            for x in range(i_idx, len(cols) + i_idx):
                arr[y, x] = _round_if_number(df_val[y - i_col, x - i_idx])
        return arr

    def _add_table(doc, arr, right_every_second=False, right_every_after=False, all_centered=False):
        table = doc.add_table(rows=(arr.shape[0]), cols=arr.shape[1])
        for y in range(arr.shape[0]):
            for x in range(arr.shape[1]):
                table.cell(y, x).text = str(arr[y, x])
                if right_every_second and x % 2 != 0:
                    table.cell(y, x).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if right_every_after != False and x >= right_every_after:
                    table.cell(y, x).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if all_centered:
                    table.cell(y, x).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    resids_df, coef_df, params = _modified_regressor_summary(model, X.values, y.values, X_cols)
    document = Document()
    document.add_heading(titel)


    model_name = str(model) if str(model).find('(') == -1 else str(model)[:str(model).find('(')]
    summ_stats = {'Dep. Variable:': y.name,
                  'Model:': model_name,
                  'No. Observations:': len(X),
                  'R-squared:': params['R2'],
                  'Adj. R-squared:': params['Adj R2'],
                  'F-statistic:': params['F-statistic']}
    for key, value in overwrite_summary.items():
        summ_stats[key] = value
    summ_stats = _dict_to_np(summ_stats)
    np_arr = summ_stats

    coef_df.rename(columns={'Estimate':'coef', 'Std. Error':'std err', 't value':'t', 'p value':'P>|t|'}, inplace=True)
    coef_arr = _df_to_np(coef_df)
    res_arr = _df_to_np(resids_df, ignore_idx=True)

    _add_table(doc=document, arr=np.array([[titel]]), all_centered=True)
    _add_table(doc=document, arr=np_arr, right_every_second=True)
    _add_table(doc=document, arr=coef_arr, right_every_after=1)
    _add_table(doc=document, arr=np.array([['Residuals']]), all_centered=True)
    _add_table(doc=document, arr=res_arr, all_centered=True)

    document.save(output_file)






def mlflow_last_run_add_param(param_dict):
    last_experiment_id, last_run_id, _ = mlflow.search_runs(order_by=['attribute.end_time DESC'])[['experiment_id', 'run_id', 'end_time']].iloc[0]

    mlflow_dict = {'layer_df': ['notes'],
                   'model_name': ['tag', 'param'],
                   'max_epochs': ['param'],
                   'actual_epochs': ['param'],
                   'early_stopped': ['param'],
                   'loss': ['param', 'tag'],
                   'data_dataset': ['tag', 'param'],
                   'data_y_col': ['tag', 'param'],
                   'data_window_input_width': ['tag', 'param'],
                   'data_window_pred_width': ['tag', 'param'],
                   'data_split_method': ['tag', 'param'],
                   'data_normalize_method': ['tag', 'param'],
                   'data_lagged_cols': ['param'],
                   'data_hash_first_level': ['tag', 'param'],
                   'data_hash_second_level': ['tag', 'param'],
                   'data_hash': ['tag', 'param'],
                   'look_ups': ['artifact'],
                   'final_data_file': ['artifact'],
                   'data_statistics': ['artifact'],
                   'data_props': [],
                   'data_time_period': ['tag', 'param'],
                   'data_train_sample_size': ['param']}

    notes = ''
    if "model_name" in param_dict:
        notes = notes + f'# {param_dict["model_name"]}\n'

    if "kwargs" in param_dict:
        for key, value in param_dict['kwargs'].items():
            param_dict[key] = value
            mlflow_dict[key] = ['param']
        param_dict.pop('kwargs')

    if "data_props" in param_dict:
        param_dict['data_dataset'] = param_dict['data_props']['first_step']['dataset']
        param_dict['data_y_col'] = param_dict['data_props']['first_step']['dataset_y_col']
        param_dict['data_window_input_width'] = param_dict['data_props']['first_step']['window_input_width']
        param_dict['data_window_pred_width'] = param_dict['data_props']['first_step']['window_pred_width']
        param_dict['data_split_method'] = param_dict['data_props']['second_step']['split_method']
        param_dict['data_normalize_method'] = param_dict['data_props']['second_step']['normalize_method']
        #param_dict['data_lagged_cols'] = str(param_dict['data_props']['second_step']['lagged_col_dict'])
        param_dict['data_hash_first_level'] = param_dict['data_props']['first_step_data_hash']
        param_dict['data_hash_second_level'] = param_dict['data_props']['second_step_data_hash']
        param_dict['data_hash'] = param_dict['data_hash_first_level'] + '_' + param_dict['data_hash_second_level']
        param_dict['data_time_period'] = param_dict['data_props']['iter_step']
        param_dict['data_train_sample_size'] = param_dict["data_props"]["statistics"]['train']['samples']

        cache_folder = get_project_directories(key='cache_dir')
        cache_folder = os.path.join(cache_folder, param_dict['data_hash_first_level'])

        look_up_file = os.path.join(cache_folder, f'{param_dict["data_hash_second_level"]}_{param_dict["data_time_period"]}_look_up.json')
        final_data_file = os.path.join(cache_folder, f'{param_dict["data_hash_second_level"]}_{param_dict["data_time_period"]}_data_schema.json')
        statistics_file = os.path.join(cache_folder, f'{param_dict["data_hash_second_level"]}_{param_dict["data_time_period"]}_data_statistics.txt')

        if not os.path.exists(statistics_file):
            with open(statistics_file, "w") as outfile:
                txt = tabulate.tabulate(pd.DataFrame(param_dict["data_props"]["statistics"]), headers='keys', tablefmt='simple')
                outfile.write(txt)

        if not os.path.exists(look_up_file):
            with open(look_up_file, "w") as outfile:
                json.dump(param_dict['data_props']['look_ups'], outfile)

        if not os.path.exists(final_data_file):
            for key, value in param_dict['data_props']['final_data']['idx'].items():
                param_dict['data_props']['final_data']['idx'][key] = value.tolist()
            with open(final_data_file, "w") as outfile:
                json.dump(param_dict['data_props']['final_data'], outfile)

        param_dict['look_ups'] = look_up_file
        param_dict['final_data_file'] = final_data_file
        param_dict['data_statistics'] = statistics_file

        notes = notes + f'**Time:** {param_dict["data_time_period"]}\n**Dataset:** {param_dict["data_dataset"]}\n'




    with mlflow.start_run(run_id=last_run_id) as run:
        for key, value in param_dict.items():
            if key[:8] == 'metrics_':
                type_list = ['metric_dict']
            else:
                type_list = mlflow_dict[key]
            for i_type in type_list:
                if i_type == 'tag':
                    mlflow.set_tag(key, value)
                elif i_type == 'param':
                    mlflow.log_param(key, value)
                elif i_type == 'notes':
                    pass
                elif i_type == 'metric':
                    mlflow.log_metric(key, value)
                elif i_type == 'metric_dict':
                    for metric_key, metric_value in value.items():
                        mlflow.log_metric(key[8:] + '_' + metric_key, metric_value)
                elif i_type == 'artifact':
                    mlflow.log_artifact(value)
                else:
                    raise Exception(f'Unknown mlflow type {i_type}')


        if "layer_df" in param_dict:
            notes = notes + (param_dict["layer_df"].to_markdown(index=False) if len(param_dict["layer_df"]) > 0 else '\nNo layers.')
        if "metrics_test" in param_dict:
            metrics_test = pd.DataFrame.from_dict({key: [value] for key, value in param_dict["metrics_test"].items()}, orient='columns')
            notes = notes + '\n\n<br>\n## Model Performance Metrics:\n' + (metrics_test.to_markdown(index=False) if len(param_dict["metrics_test"]) > 0 else 'Empty test metrics.')
        if "data_props" in param_dict:
            notes = notes + '\n\n<br>\n## Data Statistics:\n' + pd.DataFrame(param_dict["data_props"]["statistics"]).to_markdown(index=True)

        mlflow.set_tag("mlflow.note.content", notes)



def sort_list_of_sub(sub_li, sort_element=0):
    l = len(sub_li)
    for i in range(0, l):
        for j in range(0, l-i-1):
            if (sub_li[j][sort_element] > sub_li[j + 1][sort_element]):
                tempo = sub_li[j]
                sub_li[j]= sub_li[j + 1]
                sub_li[j + 1]= tempo
    return sub_li




def multiprocessing_func_with_progressbar(func, argument_list, num_processes=-1, results='append'):
    if num_processes == -1:
        num_processes = multiprocessing.cpu_count()
    pool = multiprocessing.Pool(processes=num_processes)

    jobs = [pool.apply_async(func=func, args=(*argument,)) if isinstance(argument, tuple) else pool.apply_async(func=func, args=(argument,)) for argument in argument_list]
    pool.close()

    list_a, list_b, list_c, list_d, list_e, list_f, list_g, list_h = [], [], [], [], [], [], [], []
    list_of_lists = [list_a, list_b, list_c, list_d, list_e, list_f, list_g, list_h]

    bar_format = '[elapsed: {elapsed} min] |{bar:50}| {percentage:3.0f}% - ETA: {remaining}  ({n_fmt}/{total_fmt} - {rate_fmt} - ' + str(num_processes) + ' simultaneous processes)'
    for job in tqdm(jobs, bar_format=bar_format):
        out = job.get()
        if type(out) == tuple:
            if len(out) > 8:
                raise Exception('Please modify multiprocessing_func_with_progressbar function. Function is currently just able to handle max 8 retun arguments.')
            for x, lst in zip(out, list_of_lists):
                if results == 'append':
                    lst.append(x)
                elif results == 'extend':
                    lst.extend(x)
                else:
                    raise Exception(f'Unknown input for results={results}. Valid arguments are append or extend.')
        else:
            if results == 'append':
                list_a.append(out)
            elif results == 'extend':
                list_a.extend(out)
            else:
                raise Exception(f'Unknown input for results={results}. Valid arguments are append or extend.')

    len_out = len(out) if type(out) == tuple else 1
    if len_out == 1:
        out = list_of_lists[0]
    else:
        out = tuple(list_of_lists[:len_out])

    return out









if __name__ == '__main__':
    #print(get_project_directories())
    #print(get_project_directories('project_dir'))
    #print(get_credentials())
    #print(get_credentials('reuters_eikon_api'))

    #get_credentials()
    print(test_b)
    print(sort_list_of_sub(test_b, sort_element=1))



